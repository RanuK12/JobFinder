"""
CV Parser Module

Extracts text content from PDF and DOCX files for job matching.
Supports multiple file formats and provides robust error handling.
"""

import os
import re
import logging
from typing import Optional

from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document

# Configure module logger
logger = logging.getLogger(__name__)


class CVParserError(Exception):
    """Custom exception for CV parsing errors."""
    pass


def parse_cv(filepath: str) -> Optional[str]:
    """
    Extract text from a CV file (PDF or DOCX).

    Args:
        filepath: Path to the CV file.

    Returns:
        Extracted and cleaned text content, or None if extraction fails.

    Raises:
        CVParserError: If the file format is unsupported.
    """
    if not filepath:
        logger.error("No filepath provided")
        return None

    if not os.path.exists(filepath):
        logger.error(f"File does not exist: {filepath}")
        return None

    file_ext = os.path.splitext(filepath)[1].lower()
    logger.info(f"Processing file: {filepath} (extension: {file_ext})")

    try:
        if file_ext == '.pdf':
            text = _extract_from_pdf(filepath)
        elif file_ext == '.docx':
            text = _extract_from_docx(filepath)
        else:
            logger.error(f"Unsupported file format: {file_ext}")
            return None

        if not text:
            logger.warning(f"No text content extracted from: {filepath}")
            return None

        # Clean and normalize the extracted text
        cleaned_text = _clean_text(text)

        if not cleaned_text:
            logger.warning("Text is empty after cleaning")
            return None

        logger.info(
            f"Successfully extracted {len(cleaned_text)} characters from {filepath}"
        )
        return cleaned_text

    except Exception as e:
        logger.error(f"Error processing file {filepath}: {str(e)}")
        return None


def _extract_from_pdf(filepath: str) -> Optional[str]:
    """
    Extract text from a PDF file using pdfminer.

    Args:
        filepath: Path to the PDF file.

    Returns:
        Raw text content or None.
    """
    try:
        text = pdf_extract_text(filepath)
        if text and text.strip():
            logger.info(f"PDF extraction successful: {len(text)} chars")
            return text
        logger.warning("PDF extraction returned empty text")
        return None
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        return None


def _extract_from_docx(filepath: str) -> Optional[str]:
    """
    Extract text from a DOCX file.

    Args:
        filepath: Path to the DOCX file.

    Returns:
        Raw text content or None.
    """
    try:
        doc = Document(filepath)

        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)

        if paragraphs:
            text = '\n'.join(paragraphs)
            logger.info(
                f"DOCX extraction successful: {len(paragraphs)} paragraphs"
            )
            return text

        logger.warning("DOCX extraction returned no content")
        return None
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        return None


def _clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.

    - Removes excessive whitespace
    - Normalizes line endings
    - Removes non-printable characters

    Args:
        text: Raw text to clean.

    Returns:
        Cleaned text string.
    """
    if not isinstance(text, str):
        text = str(text)

    # Remove non-printable characters (keep newlines and spaces)
    text = re.sub(r'[^\x20-\x7E\xA0-\xFF\n\t]', ' ', text)

    # Normalize multiple whitespace (but preserve single newlines for structure)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Strip leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)

    return text.strip()


def extract_sections(text: str) -> dict:
    """
    Attempt to identify common CV sections from text.

    Args:
        text: Cleaned CV text.

    Returns:
        Dictionary with identified sections.
    """
    sections = {
        'contact': '',
        'summary': '',
        'experience': '',
        'education': '',
        'skills': '',
        'languages': '',
        'other': ''
    }

    # Common section headers (English and Spanish)
    section_patterns = {
        'summary': r'(?i)(summary|profile|about|objetivo|perfil|resumen)',
        'experience': r'(?i)(experience|work|employment|experiencia|trabajo)',
        'education': r'(?i)(education|academic|estudios|educación|formación)',
        'skills': r'(?i)(skills|technologies|competencias|habilidades|tecnologías)',
        'languages': r'(?i)(languages|idiomas)',
    }

    current_section = 'other'
    lines = text.split('\n')

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # Check if this line is a section header
        for section_name, pattern in section_patterns.items():
            if re.match(pattern, line_stripped):
                current_section = section_name
                break

        sections[current_section] += line_stripped + '\n'

    return sections
